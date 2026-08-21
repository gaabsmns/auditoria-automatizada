"""
APP DE AUDITORIA AUTOMATIZADA

VERSAO RESISTENTE A AMBIENTES SEM:
- streamlit
- python-docx

Porque dependencia quebrada e praticamente patrimonio cultural da programacao.
"""

from io import BytesIO
from pathlib import Path
import argparse
import os
import re
import unicodedata
import pandas as pd

# ==================================================
# STREAMLIT OPCIONAL
# ==================================================
try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ModuleNotFoundError:
    STREAMLIT_AVAILABLE = False

# ==================================================
# PYTHON-DOCX OPCIONAL
# ==================================================
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ModuleNotFoundError:
    DOCX_AVAILABLE = False

# ==================================================
# CONFIGURACAO STREAMLIT
# ==================================================
if STREAMLIT_AVAILABLE:
    st.set_page_config(
        page_title="Auditoria Automatizada",
        layout="wide"
    )

    st.title("Auditoria Automatizada de Faturamento")
    st.caption(
        "Sistema de conferencia automatica"
    )


def obter_usuarios_app():
    usuarios = {}

    usuarios_env = os.getenv("AUDITORIA_USUARIOS")

    if usuarios_env:
        for item in usuarios_env.split(";"):
            if ":" not in item:
                continue

            usuario, senha = item.split(":", 1)
            usuarios[usuario.strip().lower()] = senha.strip()

    if STREAMLIT_AVAILABLE:
        try:
            usuarios_secrets = st.secrets.get("USUARIOS", {})

            for usuario, senha in usuarios_secrets.items():
                usuarios[str(usuario).strip().lower()] = str(senha)
        except Exception:
            pass

    return usuarios


def obter_senha_geral_app():
    senha = os.getenv("AUDITORIA_SENHA")

    if senha:
        return senha

    if STREAMLIT_AVAILABLE:
        try:
            return st.secrets.get("AUDITORIA_SENHA")
        except Exception:
            return None

    return None


def verificar_acesso_streamlit():
    usuarios_configurados = obter_usuarios_app()
    senha_geral = obter_senha_geral_app()

    if not usuarios_configurados and not senha_geral:
        st.warning(
            "Login nao configurado. Para publicar com seguranca, configure USUARIOS ou AUDITORIA_SENHA."
        )
        return True

    if st.session_state.get("acesso_liberado"):
        return True

    usuario_digitado = st.text_input(
        "Usuario/E-mail"
    ).strip().lower()

    senha_digitada = st.text_input(
        "Senha",
        type="password"
    )

    if st.button("Entrar"):
        usuario_valido = (
            usuario_digitado in usuarios_configurados and
            senha_digitada == usuarios_configurados[usuario_digitado]
        )

        senha_geral_valida = (
            senha_geral and
            senha_digitada == senha_geral and
            not usuarios_configurados
        )

        if usuario_valido or senha_geral_valida:
            st.session_state["acesso_liberado"] = True
            st.session_state["usuario_logado"] = usuario_digitado
            st.rerun()
        else:
            st.error("Usuario ou senha incorretos")

    return False

# ==================================================
# LEITURA DO ARQUIVO
# ==================================================
def texto_normalizado(valor):
    texto = "" if pd.isna(valor) else str(valor).strip()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    return texto.upper()


def normalizar_nome_coluna(coluna):
    nome = texto_normalizado(coluna)
    nome = re.sub(r"\s+", " ", nome)

    aliases = {
        "DATA": "DATA",
        "HORA": "HORA",
        "PACIENTE": "PACIENTE",
        "PROCEDIMENTO": "Procedimento",
        "PROFISSIONAL": "Profissional",
        "PROFISSIONA": "Profissional",
        "FORNECEDOR": "Fornecedor",
        "QUANTIDADE": "Quantidade",
        "QTD": "Quantidade",
        "UNIDADE BASICA": "Unidade Basica",
        "UBS": "Unidade Basica",
        "UNDIADE BASICA": "Unidade Basica",
        "LOCAL DE ATENDIMENTO": "Local de Atendimento",
    }

    return aliases.get(nome, str(coluna).strip())


def encontrar_linha_cabecalho(df_sem_cabecalho):
    melhores_colunas = {"PROCEDIMENTO", "PROFISSIONAL", "PROFISSIONA", "FORNECEDOR"}

    for indice, linha in df_sem_cabecalho.iterrows():
        valores = {texto_normalizado(valor) for valor in linha.dropna().tolist()}

        if len(valores & melhores_colunas) >= 3:
            return indice

    return None


def extrair_municipio_origem(df_sem_cabecalho):
    for _, linha in df_sem_cabecalho.head(6).iterrows():
        for valor in linha.dropna().tolist():
            texto = str(valor).strip()

            if not texto:
                continue

            texto_limpo = texto_normalizado(texto)

            if (
                "RELATORIO" in texto_limpo or
                "PERIODO" in texto_limpo or
                "CISMETRO" in texto_limpo
            ):
                continue

            return texto

    return "municipio"


def aplicar_cabecalho_detectado(df_sem_cabecalho):
    municipio = extrair_municipio_origem(df_sem_cabecalho)
    linha_cabecalho = encontrar_linha_cabecalho(df_sem_cabecalho)

    if linha_cabecalho is None:
        df_sem_cabecalho.attrs["municipio_origem"] = municipio
        return df_sem_cabecalho

    colunas = df_sem_cabecalho.iloc[linha_cabecalho].tolist()
    df = df_sem_cabecalho.iloc[linha_cabecalho + 1:].copy()
    df.columns = colunas
    df = df.dropna(how="all")
    df.attrs["municipio_origem"] = municipio

    return df


def extrair_horas_procedimento(procedimento):
    texto = texto_normalizado(procedimento)
    match = re.search(r"(\d{1,2})\s*(HORA|HORAS|HRS|HR)\b", texto)

    if match:
        return float(match.group(1))

    return pd.NA


def eh_sessao(procedimento):
    return texto_normalizado(procedimento).startswith("SESSAO")


def eh_plantao(procedimento):
    return "PLANTAO" in texto_normalizado(procedimento)


def carregar_arquivo(file):
    nome = getattr(file, "name", str(file)).lower()

    if nome.endswith(".csv"):
        df = pd.read_csv(file, header=None)
        return aplicar_cabecalho_detectado(df)

    df = pd.read_excel(file, header=None)
    return aplicar_cabecalho_detectado(df)

# ==================================================
# NORMALIZACAO
# ==================================================
def normalizar_dados(df):
    df = df.copy()

    df.columns = [normalizar_nome_coluna(c) for c in df.columns]

    obrigatorias = [
        "Procedimento",
        "Profissional",
        "Fornecedor"
    ]

    faltando = [
        c for c in obrigatorias
        if c not in df.columns
    ]

    if faltando:
        raise ValueError(
            f"Colunas obrigatorias ausentes: {', '.join(faltando)}"
        )

    if "DATA" in df.columns:
        df["DATA"] = pd.to_datetime(
            df["DATA"],
            errors="coerce"
        )

    if "HORA" in df.columns:
        df["HORA"] = (
            df["HORA"]
            .astype(str)
            .str.strip()
        )

    texto_cols = [
        "Procedimento",
        "Profissional",
        "Fornecedor"
    ]

    if "PACIENTE" in df.columns:
        texto_cols.append("PACIENTE")

    for col in texto_cols:
        df[col] = (
            df[col]
            .fillna("")
            .astype(str)
            .str.strip()
        )

    if "DATA" in df.columns and "HORA" in df.columns:
        data_texto = df["DATA"].dt.strftime("%Y-%m-%d")

        df["DATAHORA"] = pd.to_datetime(
            data_texto + " " + df["HORA"],
            errors="coerce"
        )
    else:
        df["DATAHORA"] = pd.NaT

    if "DuracaoHoras" in df.columns:
        df["DuracaoHoras"] = pd.to_numeric(
            df["DuracaoHoras"],
            errors="coerce"
        )
        df["__DuracaoHorasOrigem"] = "coluna"
    else:
        df["DuracaoHoras"] = df["Procedimento"].apply(extrair_horas_procedimento)
        df["__DuracaoHorasOrigem"] = "procedimento"

    if "Quantidade" in df.columns:
        df["Quantidade"] = pd.to_numeric(
            df["Quantidade"],
            errors="coerce"
        ).fillna(0)
    else:
        df["Quantidade"] = 1

    if "DATA" in df.columns and "HORA" in df.columns:
        return df.dropna(subset=["DATAHORA"])

    return df


def dataframe_vazio(df_base):
    return pd.DataFrame(columns=df_base.columns)


def tem_colunas(df, colunas):
    return all(coluna in df.columns for coluna in colunas)


def tem_sobreposicao_real(inicio_atual, fim_atual, inicio_anterior, fim_anterior):
    return inicio_anterior < fim_atual and inicio_atual < fim_anterior

# ==================================================
# REGRA 1
# ==================================================
def regra_1(df):
    if not tem_colunas(df, ["Fornecedor", "Profissional", "DATA", "HORA", "PACIENTE"]):
        return dataframe_vazio(df)

    base = df[
        ~df["Procedimento"].apply(eh_sessao)
    ].copy()

    if base.empty:
        return dataframe_vazio(df)

    resultado = (
        base.groupby([
            "Fornecedor",
            "Profissional",
            "DATA",
            "HORA"
        ])
        .filter(lambda x: x["PACIENTE"].nunique() > 1)
        .sort_values([
            "Profissional",
            "DATA",
            "HORA"
        ])
    )

    return resultado.reset_index(drop=True)

# ==================================================
# REGRA 2
# ==================================================
def regra_2(df):
    if not tem_colunas(df, ["Profissional", "PACIENTE", "DATAHORA"]):
        return dataframe_vazio(df)

    sessao = df[
        df["Procedimento"]
        .apply(eh_sessao)
    ].copy()

    if sessao.empty:
        return dataframe_vazio(df)

    conflitos = []

    for _, grupo in sessao.groupby([
        "Profissional",
        "PACIENTE"
    ]):

        grupo = grupo.sort_values("DATAHORA")

        for i in range(len(grupo) - 1):
            atual = grupo.iloc[i]
            proximo = grupo.iloc[i + 1]

            intervalo = (
                proximo["DATAHORA"] - atual["DATAHORA"]
            ).total_seconds() / 60

            if 0 <= intervalo < 30:
                conflitos.extend([atual, proximo])

    if conflitos:
        return pd.DataFrame(conflitos).drop_duplicates().reset_index(drop=True)

    return dataframe_vazio(df)

# ==================================================
# REGRA 3
# ==================================================
def regra_3(df):
    if not tem_colunas(df, ["Profissional", "Procedimento", "DATAHORA", "DuracaoHoras"]):
        return dataframe_vazio(df)

    plantoes = df[
        df["Procedimento"]
        .apply(eh_plantao)
    ].copy()

    if plantoes.empty:
        return dataframe_vazio(df)

    plantoes = plantoes.dropna(subset=["DuracaoHoras"])

    if plantoes.empty:
        return dataframe_vazio(df)

    plantoes["INICIO"] = plantoes["DATAHORA"]

    plantoes["FIM"] = (
        plantoes["INICIO"] +
        pd.to_timedelta(plantoes["DuracaoHoras"], unit="h")
    )

    conflitos = []

    for _, grupo in plantoes.groupby("Profissional"):
        grupo = grupo.sort_values("INICIO")

        for i in range(len(grupo) - 1):
            atual = grupo.iloc[i]
            proximo = grupo.iloc[i + 1]

            if tem_sobreposicao_real(
                proximo["INICIO"],
                proximo["FIM"],
                atual["INICIO"],
                atual["FIM"]
            ):
                conflitos.extend([atual, proximo])

    if conflitos:
        return pd.DataFrame(conflitos).drop_duplicates().reset_index(drop=True)

    return dataframe_vazio(df)

# ==================================================
# REGRA 4
# ==================================================
def regra_4(df):
    if not tem_colunas(df, ["PACIENTE", "Profissional", "DATAHORA"]):
        return dataframe_vazio(df)

    sessao = df[
        df["Procedimento"]
        .apply(eh_sessao)
    ].copy()

    if sessao.empty:
        return dataframe_vazio(df)

    conflitos = []

    for _, grupo in sessao.groupby("PACIENTE"):
        grupo = grupo.sort_values("DATAHORA")

        for i in range(len(grupo) - 1):
            atual = grupo.iloc[i]
            proximo = grupo.iloc[i + 1]

            intervalo = (
                proximo["DATAHORA"] - atual["DATAHORA"]
            ).total_seconds() / 60

            if (
                atual["Profissional"] != proximo["Profissional"] and
                0 <= intervalo < 30
            ):
                conflitos.extend([atual, proximo])

    if conflitos:
        return pd.DataFrame(conflitos).drop_duplicates().reset_index(drop=True)

    return dataframe_vazio(df)

# ==================================================
# REGRA 5
# ==================================================
def regra_5(df):
    if not tem_colunas(df, ["Fornecedor", "Profissional", "PACIENTE", "Procedimento", "DATA"]):
        return dataframe_vazio(df)

    base = df.copy()
    base = base[
        ~base["Procedimento"].apply(eh_sessao)
    ].copy()

    if base.empty:
        return dataframe_vazio(df)

    base = base[
        base["PACIENTE"].astype(str).str.strip().ne("") &
        base["Procedimento"].astype(str).str.strip().ne("")
    ].copy()

    if base.empty:
        return dataframe_vazio(df)

    base["__FORNECEDOR_DUP"] = base["Fornecedor"].apply(texto_normalizado)
    base["__PROFISSIONAL_DUP"] = base["Profissional"].apply(texto_normalizado)
    base["__PACIENTE_DUP"] = base["PACIENTE"].apply(texto_normalizado)
    base["__PROCEDIMENTO_DUP"] = base["Procedimento"].apply(texto_normalizado)
    base["__DATA_DUP"] = base["DATA"].dt.strftime("%Y-%m-%d")

    chaves = [
        "__FORNECEDOR_DUP",
        "__PROFISSIONAL_DUP",
        "__PACIENTE_DUP",
        "__PROCEDIMENTO_DUP",
        "__DATA_DUP"
    ]

    base["QtdDuplicidade"] = (
        base
        .groupby(chaves, dropna=False)["PACIENTE"]
        .transform("size")
    )

    resultado = base[
        base["QtdDuplicidade"] > 1
    ].sort_values([
        "PACIENTE",
        "Procedimento",
        "DATA",
        "HORA"
    ])

    return resultado.reset_index(drop=True)

# ==================================================
# REGRA 6
# ==================================================
def regra_6(df):
    if not tem_colunas(df, ["Profissional", "Procedimento", "DATAHORA", "DuracaoHoras"]):
        return dataframe_vazio(df)

    plantoes = df[
        df["Procedimento"]
        .apply(eh_plantao)
    ].copy()

    outros = df[
        ~df["Procedimento"].apply(eh_plantao)
    ].copy()

    if plantoes.empty:
        return dataframe_vazio(df)

    plantoes = plantoes.dropna(subset=["DuracaoHoras"])

    if plantoes.empty:
        return dataframe_vazio(df)

    plantoes["INICIO"] = plantoes["DATAHORA"]

    plantoes["FIM"] = (
        plantoes["INICIO"] +
        pd.to_timedelta(plantoes["DuracaoHoras"], unit="h")
    )

    conflitos = []

    for _, plantao in plantoes.iterrows():
        conflito = outros[
            (outros["Profissional"] == plantao["Profissional"]) &
            (outros["DATAHORA"] >= plantao["INICIO"]) &
            (outros["DATAHORA"] < plantao["FIM"])
        ]

        if not conflito.empty:
            conflitos.append(conflito)

    if conflitos:
        return pd.concat(conflitos).drop_duplicates().reset_index(drop=True)

    return dataframe_vazio(df)

# ==================================================
# REGRA 7
# ==================================================
def regra_7(df):
    if "DuracaoHoras" not in df.columns:
        return dataframe_vazio(df)

    if (
        "__DuracaoHorasOrigem" in df.columns and
        not (df["__DuracaoHorasOrigem"] == "coluna").any()
    ):
        return dataframe_vazio(df)

    plantoes = df[
        df["Procedimento"]
        .apply(eh_plantao)
    ].copy()

    if "__DuracaoHorasOrigem" in plantoes.columns:
        plantoes = plantoes[
            plantoes["__DuracaoHorasOrigem"] == "coluna"
        ].copy()

    if plantoes.empty:
        return dataframe_vazio(df)

    conflitos = []

    for _, row in plantoes.iterrows():
        horas = row["DuracaoHoras"]
        quantidade = row.get("Quantidade", 0)

        if pd.isna(horas):
            conflitos.append(row)
            continue

        minimo = 4 if horas >= 4 else int(horas)

        if quantidade < minimo:
            conflitos.append(row)

    if conflitos:
        return pd.DataFrame(conflitos).reset_index(drop=True)

    return dataframe_vazio(df)


# ==================================================
# REGRA 8
# ==================================================
def eh_plantao_enfermagem(procedimento):
    texto = texto_normalizado(procedimento)

    if "PLANTAO" not in texto:
        return False

    return (
        "PLANTAO ENFERMEIRO" in texto or
        "PLANTAO TECNICO DE ENFERMAGEM" in texto
    )


def categoria_enfermagem(procedimento):
    texto = texto_normalizado(procedimento)

    if "TECNICO DE ENFERMAGEM" in texto:
        return "Tecnico de enfermagem"

    if "ENFERMEIRO" in texto:
        return "Enfermeiro"

    return ""


def eh_plantao_raio_x(procedimento):
    texto = texto_normalizado(procedimento)
    return "PLANTAO DE TECNICO DE RAIO X" in texto


def inicio_semana_domingo(data):
    data = pd.to_datetime(data).normalize()
    dias_desde_domingo = (data.weekday() + 1) % 7
    return data - pd.Timedelta(days=dias_desde_domingo)


def regra_8(df):
    if not tem_colunas(df, ["Fornecedor", "Profissional", "Procedimento", "DuracaoHoras", "Quantidade"]):
        return dataframe_vazio(df)

    apontamentos = []
    base = df.copy()
    base["HorasCalculadas"] = base["DuracaoHoras"] * base["Quantidade"]

    enfermagem = base[
        base["Procedimento"].apply(eh_plantao_enfermagem) &
        base["HorasCalculadas"].notna()
    ].copy()

    if not enfermagem.empty:
        enfermagem["Categoria"] = enfermagem["Procedimento"].apply(categoria_enfermagem)

        if "DATA" in enfermagem.columns and enfermagem["DATA"].notna().any():
            enfermagem["MesReferencia"] = enfermagem["DATA"].dt.to_period("M").astype(str)
            grupos_mensais = [
                "Fornecedor",
                "Profissional",
                "Categoria",
                "MesReferencia"
            ]
        else:
            enfermagem["MesReferencia"] = "Periodo da planilha"
            grupos_mensais = [
                "Fornecedor",
                "Profissional",
                "Categoria",
                "MesReferencia"
            ]

        total_mensal = (
            enfermagem
            .groupby(grupos_mensais, dropna=False)["HorasCalculadas"]
            .sum()
            .reset_index()
        )

        total_mensal = total_mensal[total_mensal["HorasCalculadas"] > 180]

        for _, row in total_mensal.iterrows():
            apontamentos.append({
                "TipoApontamento": "Enfermagem acima de 180h mensais",
                "Fornecedor": row["Fornecedor"],
                "Profissional": row["Profissional"],
                "Categoria": row["Categoria"],
                "MesReferencia": row["MesReferencia"],
                "TotalHoras": row["HorasCalculadas"],
                "LimiteHoras": 180,
                "ExcessoHoras": row["HorasCalculadas"] - 180,
            })

    raio_x = base[
        base["Procedimento"].apply(eh_plantao_raio_x) &
        base["HorasCalculadas"].notna()
    ].copy()

    if (
        not raio_x.empty and
        "DATA" in raio_x.columns and
        raio_x["DATA"].notna().any()
    ):
        diario = (
            raio_x
            .groupby(["Fornecedor", "Profissional", "DATA"], dropna=False)["HorasCalculadas"]
            .sum()
            .reset_index()
        )

        diario = diario[diario["HorasCalculadas"] > 6]

        for _, row in diario.iterrows():
            apontamentos.append({
                "TipoApontamento": "Raio X acima de 6h no dia",
                "Fornecedor": row["Fornecedor"],
                "Profissional": row["Profissional"],
                "DATA": row["DATA"],
                "TotalHoras": row["HorasCalculadas"],
                "LimiteHoras": 6,
                "ExcessoHoras": row["HorasCalculadas"] - 6,
            })

        raio_x["SemanaInicio"] = raio_x["DATA"].apply(inicio_semana_domingo)
        raio_x["SemanaFim"] = raio_x["SemanaInicio"] + pd.Timedelta(days=6)

        semanal = (
            raio_x
            .groupby(["Fornecedor", "Profissional", "SemanaInicio", "SemanaFim"], dropna=False)["HorasCalculadas"]
            .sum()
            .reset_index()
        )

        semanal = semanal[semanal["HorasCalculadas"] > 24]

        for _, row in semanal.iterrows():
            apontamentos.append({
                "TipoApontamento": "Raio X acima de 24h semanais",
                "Fornecedor": row["Fornecedor"],
                "Profissional": row["Profissional"],
                "SemanaInicio": row["SemanaInicio"],
                "SemanaFim": row["SemanaFim"],
                "TotalHoras": row["HorasCalculadas"],
                "LimiteHoras": 24,
                "ExcessoHoras": row["HorasCalculadas"] - 24,
            })

    if apontamentos:
        return pd.DataFrame(apontamentos).reset_index(drop=True)

    return pd.DataFrame(columns=[
        "TipoApontamento",
        "Fornecedor",
        "Profissional",
        "Categoria",
        "MesReferencia",
        "DATA",
        "SemanaInicio",
        "SemanaFim",
        "TotalHoras",
        "LimiteHoras",
        "ExcessoHoras",
    ])

# ==================================================
# GERACAO RELATORIO
# ==================================================
def gerar_relatorio_texto(resultados):
    linhas = []

    linhas.append("APONTAMENTOS DE AUDITORIA")
    linhas.append("=" * 50)
    linhas.append("")

    for nome_regra, df_regra in resultados.items():
        linhas.append(nome_regra)
        linhas.append("-" * 50)

        if df_regra.empty:
            linhas.append("Nenhuma ocorrencia encontrada")
            linhas.append("")
            continue

        for _, row in df_regra.iterrows():
            texto = (
                f"{row.get('DATA', '')} | "
                f"{row.get('HORA', '')} | "
                f"Profissional: {row.get('Profissional', '')} | "
                f"Paciente: {row.get('PACIENTE', '')} | "
                f"Procedimento: {row.get('Procedimento', '')}"
            )

            linhas.append(texto)

        linhas.append("")

    return "\n".join(linhas)

# ==================================================
# GERACAO WORD
# ==================================================
FRASES_APONTAMENTOS = {
    "Regra 1": (
        "Nao pode haver dois pacientes diferentes em um mesmo horario, porem "
        "pode haver duplicidade de horario para um mesmo paciente quando realizou "
        "mais de um procedimento. Necessario verificar os lancamentos abaixo:"
    ),
    "Regra 2": (
        "Reforcamos que nas sessoes pode haver mais de um paciente por horario, "
        "mas quando um mesmo paciente tem mais de uma sessao por dia, e necessario "
        "lanca-las em horarios diferentes, com intervalo minimo de 30 minutos entre elas."
    ),
    "Regra 3": (
        "Duplicidade de horario. E necessario finalizar a carga horaria de um plantao "
        "antes de dar inicio a outro. Necessaria alteracao:"
    ),
    "Regra 4": (
        "Para pacientes com duas sessoes no mesmo dia, se faz necessario o lancamento "
        "dessas sessoes em horarios diferentes, com intervalo minimo de 30 minutos."
    ),
    "Regra 5": (
        "Verificar se nao houve duplicidade de lancamento nos registros abaixo. "
        "Caso tenha ocorrido duplicidade, e necessario excluir um dos lancamentos:"
    ),
    "Regra 6": (
        "Durante carga horaria de plantao nao devem ser lancados outros procedimentos, "
        "sendo necessario o termino da carga horaria do plantao para realizar qualquer "
        "outro lancamento. Necessario verificar:"
    ),
    "Regra 7": (
        "Necessario incluir o nome de pelo menos 4 pacientes atendidos no plantao "
        "ou justificar o lancamento com quantidade menor:"
    ),
    "Regra 8": (
        "Enviar justificativa para o cumprimento de horas excedentes no periodo pelos "
        "profissionais listados abaixo, considerando os limites de carga horaria definidos:"
    ),
}


def frase_apontamento(nome_regra):
    for prefixo, frase in FRASES_APONTAMENTOS.items():
        if nome_regra.startswith(prefixo):
            return frase

    return "Verificar os lancamentos relacionados abaixo:"


def colunas_word_para_regra(df, nome_regra):
    df_saida = preparar_saida_apontamentos(df, nome_regra)

    preferencias = [
        "DATA",
        "HORA",
        "PACIENTE",
        "Dt. Nasc.",
        "Procedimento",
        "Quantidade",
        "Profissional",
        "Fornecedor",
        "Categoria",
        "MesReferencia",
        "SemanaInicio",
        "SemanaFim",
        "TotalHoras",
        "LimiteHoras",
        "ExcessoHoras",
        "QtdDuplicidade",
    ]

    colunas = [coluna for coluna in preferencias if coluna in df_saida.columns]

    if not colunas:
        colunas = list(df_saida.columns[:8])

    return df_saida[colunas]


def adicionar_tabela_word(doc, df_tabela):
    if df_tabela.empty:
        return

    tabela = doc.add_table(rows=1, cols=len(df_tabela.columns))
    tabela.style = "Table Grid"

    cabecalho = tabela.rows[0].cells

    for indice, coluna in enumerate(df_tabela.columns):
        cabecalho[indice].text = str(coluna)

    for _, row in df_tabela.iterrows():
        celulas = tabela.add_row().cells

        for indice, coluna in enumerate(df_tabela.columns):
            valor = row.get(coluna, "")
            celulas[indice].text = "" if pd.isna(valor) else str(valor)


def gerar_word(resultados):

    if not DOCX_AVAILABLE:
        return gerar_relatorio_texto(resultados).encode("utf-8")

    doc = Document()

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(10)

    doc.add_heading(
        "APONTAMENTOS",
        level=1
    )

    regras_com_resultado = {
        nome_regra: df_regra
        for nome_regra, df_regra in resultados.items()
        if not df_regra.empty
    }

    if not regras_com_resultado:
        doc.add_paragraph("Nenhuma ocorrencia encontrada.")
    else:
        fornecedores = sorted({
            str(row.get("Fornecedor", "Sem fornecedor")).strip() or "Sem fornecedor"
            for df_regra in regras_com_resultado.values()
            for _, row in df_regra.iterrows()
        })

        for numero_fornecedor, fornecedor in enumerate(fornecedores, start=1):
            doc.add_heading(f"{numero_fornecedor}. {fornecedor}", level=2)

            profissionais_fornecedor = set()

            for df_regra in regras_com_resultado.values():
                df_fornecedor = df_regra[
                    df_regra["Fornecedor"].astype(str).str.strip().fillna("") == fornecedor
                ].copy()

                if df_fornecedor.empty:
                    continue

                if "Profissional" in df_fornecedor.columns:
                    profissionais_fornecedor.update(
                        df_fornecedor["Profissional"]
                        .fillna("Sem profissional")
                        .astype(str)
                        .str.strip()
                        .replace("", "Sem profissional")
                        .unique()
                    )
                else:
                    profissionais_fornecedor.add("Sem profissional")

            for profissional in sorted(profissionais_fornecedor):
                doc.add_paragraph(f"Profissional: {profissional}")

                for nome_regra, df_regra in regras_com_resultado.items():
                    df_fornecedor = df_regra[
                        df_regra["Fornecedor"].astype(str).str.strip().fillna("") == fornecedor
                    ].copy()

                    if df_fornecedor.empty:
                        continue

                    if "Profissional" in df_fornecedor.columns:
                        df_profissional = df_fornecedor[
                            df_fornecedor["Profissional"].astype(str).str.strip().fillna("") == profissional
                        ].copy()
                    else:
                        df_profissional = df_fornecedor.copy()

                    if df_profissional.empty:
                        continue

                    doc.add_paragraph(frase_apontamento(nome_regra))
                    df_tabela = colunas_word_para_regra(df_profissional, nome_regra)
                    adicionar_tabela_word(doc, df_tabela)
                    doc.add_paragraph("")

    arquivo = BytesIO()
    doc.save(arquivo)
    arquivo.seek(0)

    return arquivo


# ==================================================
# FORMATACAO DE SAIDA
# ==================================================
def preparar_saida_apontamentos(df, nome_regra=None):
    df_saida = df.copy()

    colunas_internas = [
        "DATAHORA"
    ]

    regras_com_colunas_resumidas = [
        "Regra 1",
        "Regra 3",
        "Regra 4",
        "Regra 5",
        "Regra 6",
        "Regra 7",
    ]

    colunas_para_remover = [
        "Cód.Procedimento",
        "Cod.Procedimento",
        "Vlr.Unit.R$",
        "Vlr.Total R$",
        "Unidade Basica",
        "Local de Atendimento",
        "DuracaoHoras",
    ]

    if nome_regra and any(nome_regra.startswith(regra) for regra in regras_com_colunas_resumidas):
        df_saida = df_saida.drop(
            columns=[c for c in colunas_para_remover if c in df_saida.columns],
            errors="ignore"
        )

    df_saida = df_saida.drop(
        columns=[
            c for c in df_saida.columns
            if c in colunas_internas or c.startswith("__")
        ],
        errors="ignore"
    )

    for coluna in df_saida.columns:
        if pd.api.types.is_datetime64_any_dtype(df_saida[coluna]):
            if coluna in ["DATA", "Dt. Nasc."] or coluna.startswith("Semana"):
                df_saida[coluna] = df_saida[coluna].dt.strftime("%d/%m/%Y")
            else:
                df_saida[coluna] = df_saida[coluna].dt.strftime(
                    "%d/%m/%Y %H:%M:%S"
                )

    if "Dt. Nasc." in df_saida.columns:
        data_nasc = pd.to_datetime(
            df_saida["Dt. Nasc."],
            errors="coerce",
            dayfirst=True
        )
        df_saida["Dt. Nasc."] = data_nasc.dt.strftime("%d/%m/%Y").fillna(
            df_saida["Dt. Nasc."].astype(str)
        )

    if "MesReferencia" in df_saida.columns:
        mes = pd.to_datetime(
            df_saida["MesReferencia"].astype(str) + "-01",
            errors="coerce"
        )
        df_saida["MesReferencia"] = mes.dt.strftime("%m/%Y").fillna(
            df_saida["MesReferencia"].astype(str)
        )

    return df_saida


# ==================================================
# GERACAO EXCEL
# ==================================================
def _nome_aba_excel(nome):
    proibidos = ["\\", "/", "*", "[", "]", ":", "?"]

    for caractere in proibidos:
        nome = nome.replace(caractere, "-")

    return nome[:31]


def gerar_excel(resultados):
    arquivo = BytesIO()

    resumo = pd.DataFrame([
        {
            "Tipo de apontamento": nome_regra,
            "Quantidade": len(df_regra)
        }
        for nome_regra, df_regra in resultados.items()
    ])

    with pd.ExcelWriter(arquivo, engine="openpyxl") as writer:
        resumo.to_excel(
            writer,
            sheet_name="Resumo",
            index=False
        )

        for nome_regra, df_regra in resultados.items():
            nome_aba = _nome_aba_excel(nome_regra)

            if df_regra.empty:
                pd.DataFrame({
                    "Resultado": ["Nenhuma ocorrencia encontrada"]
                }).to_excel(
                    writer,
                    sheet_name=nome_aba,
                    index=False
                )
            else:
                df_saida = preparar_saida_apontamentos(df_regra, nome_regra)

                df_saida.to_excel(
                    writer,
                    sheet_name=nome_aba,
                    index=False
                )

        for sheet in writer.book.worksheets:
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions

            for column_cells in sheet.columns:
                maior = 0
                letra = column_cells[0].column_letter

                for cell in column_cells:
                    valor = "" if cell.value is None else str(cell.value)
                    maior = max(maior, len(valor))

                sheet.column_dimensions[letra].width = min(max(maior + 2, 12), 45)

    arquivo.seek(0)

    return arquivo


def limpar_nome_arquivo(texto):
    texto = str(texto).strip()
    texto = re.sub(r'[\\/:*?"<>|]+', "-", texto)
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip(" .-") or "municipio"


def nome_base_relatorio(df):
    municipio = limpar_nome_arquivo(
        df.attrs.get("municipio_origem", "municipio")
    )
    return f"auditoria_apontamentos - {municipio}"

# ==================================================
# EXECUCAO AUDITORIA
# ==================================================
def executar_auditoria(df):
    return {
        "Regra 1 - Mesmo horario": regra_1(df),
        "Regra 2 - Sessao < 30 min": regra_2(df),
        "Regra 3 - Sobreposicao plantao": regra_3(df),
        "Regra 4 - Sessao profissionais diferentes": regra_4(df),
        "Regra 5 - Duplicidade": regra_5(df),
        "Regra 6 - Plantonista ocupado": regra_6(df),
        "Regra 7 - Validade plantao": regra_7(df),
        "Regra 8 - Limite de horas": regra_8(df)
    }

# ==================================================
# INTERFACE STREAMLIT
# ==================================================
if STREAMLIT_AVAILABLE:

    if not verificar_acesso_streamlit():
        st.stop()

    uploaded_file = st.file_uploader(
        "Envie o Excel",
        type=["xlsx", "xls", "csv"]
    )

    if uploaded_file:

        try:
            df = carregar_arquivo(uploaded_file)
            df = normalizar_dados(df)

            resultados = executar_auditoria(df)

            st.success("Arquivo processado")

            cols = st.columns(len(resultados))

            for i, (nome, resultado) in enumerate(resultados.items()):
                cols[i].metric(f"R{i+1}", len(resultado))

            for nome, resultado in resultados.items():
                with st.expander(nome):
                    if resultado.empty:
                        st.info("Nenhuma ocorrencia")
                    else:
                        st.dataframe(preparar_saida_apontamentos(resultado, nome))

            relatorio = gerar_word(resultados)
            relatorio_excel = gerar_excel(resultados)
            nome_relatorio = nome_base_relatorio(df)

            st.download_button(
                label="Baixar Excel com apontamentos",
                data=relatorio_excel,
                file_name=f"{nome_relatorio}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            if DOCX_AVAILABLE:
                st.download_button(
                    label="Baixar Word",
                    data=relatorio,
                    file_name=f"{nome_relatorio}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.download_button(
                    label="Baixar TXT",
                    data=relatorio,
                    file_name=f"{nome_relatorio}.txt",
                    mime="text/plain"
                )

        except Exception as e:
            st.error(str(e))

def mostrar_dependencias_ausentes():
    print("=" * 60)
    print("AUDITORIA AUTOMATIZADA")
    print("=" * 60)
    print()

    print("Dependencias ausentes:")

    if not STREAMLIT_AVAILABLE:
        print("- streamlit")

    if not DOCX_AVAILABLE:
        print("- python-docx")

    print()
    print("Instale com:")
    print("pip install streamlit python-docx")
    print()


def imprimir_resumo(resultados):
    print("Resumo dos apontamentos:")
    print("-" * 60)

    total = 0

    for nome, resultado in resultados.items():
        quantidade = len(resultado)
        total += quantidade
        print(f"{nome}: {quantidade}")

    print("-" * 60)
    print(f"Total de linhas apontadas: {total}")


def executar_terminal(caminho_entrada, caminho_saida=None):
    arquivo_entrada = Path(caminho_entrada)

    if not arquivo_entrada.exists():
        raise FileNotFoundError(f"Arquivo nao encontrado: {arquivo_entrada}")

    df = carregar_arquivo(arquivo_entrada)
    df = normalizar_dados(df)
    resultados = executar_auditoria(df)

    imprimir_resumo(resultados)

    if caminho_saida is None:
        nome_relatorio = nome_base_relatorio(df)
        caminho_saida = arquivo_entrada.with_name(
            f"{nome_relatorio}.txt"
        )

    caminho_saida = Path(caminho_saida)
    caminho_saida.write_text(
        gerar_relatorio_texto(resultados),
        encoding="utf-8"
    )

    caminho_excel = caminho_saida.with_suffix(".xlsx")
    relatorio_excel = gerar_excel(resultados)
    caminho_excel.write_bytes(relatorio_excel.getvalue())

    print()
    print(f"Relatorio TXT salvo em: {caminho_saida}")
    print(f"Relatorio Excel salvo em: {caminho_excel}")


def main():
    parser = argparse.ArgumentParser(
        description="Executa auditoria automatizada em arquivo CSV ou Excel."
    )
    parser.add_argument(
        "arquivo",
        nargs="?",
        help="Caminho do arquivo .csv, .xlsx ou .xls"
    )
    parser.add_argument(
        "--saida",
        help="Caminho do relatorio .txt de saida"
    )
    parser.add_argument(
        "--testes",
        action="store_true",
        help="Executa os testes internos"
    )

    args = parser.parse_args()

    if args.testes:
        _teste_regra_1()
        _teste_regra_1_ignora_sessao()
        _teste_regra_2()
        _teste_regra_5()
        _teste_regra_5_ignora_sessao()
        _teste_regra_6_plantao_sequencial_nao_conflita()
        _teste_sem_conflito()
        print("Testes internos executados com sucesso.")
        return

    if args.arquivo:
        executar_terminal(args.arquivo, args.saida)
        return

    if not STREAMLIT_AVAILABLE:
        mostrar_dependencias_ausentes()
        print("Para usar pelo terminal:")
        print("python app_auditoria.py caminho\\da\\planilha.xlsx")
        print()

# ==================================================
# TESTES
# ==================================================
def _teste_regra_1():
    dados = pd.DataFrame({
        "Fornecedor": ["X", "X"],
        "Profissional": ["ANA", "ANA"],
        "DATA": ["2025-01-01", "2025-01-01"],
        "HORA": ["08:00", "08:00"],
        "PACIENTE": ["JOAO", "MARIA"],
        "Procedimento": ["RX", "RX"]
    })

    dados = normalizar_dados(dados)

    resultado = regra_1(dados)

    assert len(resultado) == 2


def _teste_regra_1_ignora_sessao():
    dados = pd.DataFrame({
        "Fornecedor": ["X", "X"],
        "Profissional": ["ANA", "ANA"],
        "DATA": ["2025-01-01", "2025-01-01"],
        "HORA": ["08:00", "08:00"],
        "PACIENTE": ["JOAO", "MARIA"],
        "Procedimento": ["SESSÃO FISIO", "SESSÃO FISIO"]
    })

    dados = normalizar_dados(dados)

    resultado = regra_1(dados)

    assert resultado.empty


def _teste_regra_2():
    dados = pd.DataFrame({
        "Fornecedor": ["X", "X"],
        "Profissional": ["ANA", "ANA"],
        "DATA": ["2025-01-01", "2025-01-01"],
        "HORA": ["08:00", "08:20"],
        "PACIENTE": ["JOAO", "JOAO"],
        "Procedimento": ["SESSAO FISIO", "SESSAO FISIO"]
    })

    dados = normalizar_dados(dados)

    resultado = regra_2(dados)

    assert len(resultado) == 2


def _teste_regra_5():
    dados = pd.DataFrame({
        "Fornecedor": ["X", "X"],
        "Profissional": ["ANA", "ANA"],
        "DATA": ["2025-01-01", "2025-01-01"],
        "HORA": ["08:00", "09:00"],
        "PACIENTE": ["JOAO", "JOAO"],
        "Procedimento": ["RX", "RX"]
    })

    dados = normalizar_dados(dados)

    resultado = regra_5(dados)

    assert len(resultado) == 2


def _teste_regra_5_ignora_sessao():
    dados = pd.DataFrame({
        "Fornecedor": ["X", "X"],
        "Profissional": ["ANA", "ANA"],
        "DATA": ["2025-01-01", "2025-01-01"],
        "HORA": ["08:00", "09:00"],
        "PACIENTE": ["JOAO", "JOAO"],
        "Procedimento": ["SESSÃO FISIO", "SESSÃO FISIO"]
    })

    dados = normalizar_dados(dados)

    resultado = regra_5(dados)

    assert resultado.empty


def _teste_regra_6_plantao_sequencial_nao_conflita():
    dados = pd.DataFrame({
        "Fornecedor": ["X", "X"],
        "Profissional": ["WESLEY", "WESLEY"],
        "DATA": ["2026-04-10", "2026-04-10"],
        "HORA": ["07:00", "11:00"],
        "PACIENTE": ["KIRIA", "IVONE"],
        "Procedimento": [
            "PLANTÃO ENFERMEIRO - 04 Hrs",
            "PLANTÃO ENFERMEIRO - 06 HRS"
        ]
    })

    dados = normalizar_dados(dados)

    resultado_regra_3 = regra_3(dados)
    resultado_regra_6 = regra_6(dados)

    assert resultado_regra_3.empty
    assert resultado_regra_6.empty


def _teste_sem_conflito():
    dados = pd.DataFrame({
        "Fornecedor": ["X"],
        "Profissional": ["ANA"],
        "DATA": ["2025-01-01"],
        "HORA": ["08:00"],
        "PACIENTE": ["JOAO"],
        "Procedimento": ["RX"]
    })

    dados = normalizar_dados(dados)

    resultado = regra_1(dados)

    assert resultado.empty


if __name__ == "__main__":
    main()
